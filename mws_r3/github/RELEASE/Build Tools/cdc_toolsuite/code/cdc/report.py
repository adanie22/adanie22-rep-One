# -*- coding: utf-8 -*-
# 
# Computer Sciences Corporation
#
# MWS SAS Deployment Code
#
# Author: dhatchett2
# 
###########################################################################  
# Change log:
#  Version     	Date		Author	
#    1.2	  07/07/2015    A.Daniels
#   Modify setup of headers in create_xls_compute:    	
#   Remove id, image, instanceId. Replace parent with VLAN.    	
#     
#    1.1	  03/07/2015    A.Daniels
#   Modify setup of headers in create_xls_compute   	
#     
###########################################################################  
#     
#

from pprint import pprint as pp
import os
import requests
requests.packages.urllib3.disable_warnings()

import default
import re
from datetime import datetime
from lxml import etree as et
from itertools import groupby
from util import *
import xls

from docx import Document
from docx.shared import Inches
import blueprints as bps
#doc/rest/core/v3.2

def map_ver(m):
    try:
        return {
            -1: 'IN PROGRESS',
            0: 'HEAD'
        }[m]
    except KeyError:
        return m

def list_assets(asset_type,env=default.ENV,hdrs={'accept':'application/json'}):
    """
    asset_type may be "project", "environment", "container"

    r.json():

        {u'assets': [{u'checkoutAllowed': False,
                      u'id': 2,
                      u'name': u'Agility Factory',
                      u'parent': {u'href': u'https://20.148.76.21:8443/agility/api/v3.2/container/1',
                                  u'id': 1,
                                  u'name': u'Root',
                                  u'position': 0,
                                  u'rel': u'up',
                                  u'type': u'application/com.servicemesh.agility.api.Container+xml'},
                      u'removable': False,
                      u'top': False},
                     {u'checkoutAllowed': False,
                      u'id': 3,
                      u'name': u'Agility Platform',
                      u'parent': {u'href': u'https://20.148.76.21:8443/agility/api/v3.2/container/1',
                                  u'id': 1,

    """
    cfg = default.CFG[env]
    url,user,passwd = '%s/%s/search?fields=id,name,parent'%(cfg['BASE_URL'],asset_type),'%s'%cfg['USER'],'%s'%cfg['PASSWD']
    r = requests.get(url, auth=('%s'%user, '%s'%passwd),verify=False,headers=hdrs)

    if r.status_code == 200:
        for d in r.json()['assets']:
            print d['id'],d['parent']['name'],d['name']
    else:
        print 'problem listing items for asset type %s'%asset_type
        print r.status_code, url

def list_compute(asset_type,env=default.ENV,hdrs={'accept':'application/json'}):
    """
        FYI, compute keys: [u'assetPath', u'assetProperties', u'assetType', u'cloud', u'credentials', u'description', u'detailedAssetPath', u'environment', u'hostname', u'id', u'image', u'instanceId', u'lastStartedOrStoppedBy', u'lifecycleVersion', u'location', u'model', u'name', u'onboarded', u'privateAddress', u'removable', u'resources', u'scriptstatusLinks', u'stack', u'startTime', u'state', u'template', u'top', u'uuid'] 
    """
    cfg = default.CFG[env]
    url,user,passwd = '%s/compute'%(cfg['BASE_URL'],),'%s'%cfg['USER'],'%s'%cfg['PASSWD']
    r = requests.get(url, auth=('%s'%user, '%s'%passwd),verify=False,headers=hdrs)

    def set_key(d,d2,k,k2=None):
        if d2.has_key(k):
            if k2 is None:
                d.setdefault(k,d2[k])
            else:
                d.setdefault(k,d2[k][k2])
        else:
            d.setdefault(k,'')

    def set_key_b(d,d2,k,k2=None):
        if d2.has_key(k):
            if k2 is None:
                if k == 'resources':
                     for r in d2['resources']:
                            if r.has_key('network'):
                                      d.setdefault(k,r['network']['name'])
                else:
                    if k == "template":
                        d.setdefault(k,d2[k]['name'])
                    else:
                        d.setdefault(k,d2[k])					 
            else:
                d.setdefault(k,d2[k][k2])
        else:
            d.setdefault(k,'')


    if r.status_code == 200:
        l = []
        for d in r.json()['links']:
            url,user,passwd = '%s/compute/%s'%(cfg['BASE_URL'],d['id']),'%s'%cfg['USER'],'%s'%cfg['PASSWD']

            r = requests.get(url, auth=('%s'%user, '%s'%passwd),verify=False,headers=hdrs)

            j2 = r.json()
            #print(j2)

            """
            if j2.has_key('resources'):
				for r in j2['resources']:
					if r.has_key('network'):
						pp(r['network']['name'])
            if j2.has_key('state'):
				pp(j2['state'])
            if j2.has_key('template'):
				pp(j2['template']['name'])
            """

	    d2 = {}
            #orig:[set_key(d2,j2,m) for m in ('id','name','parent','instanceId','privateAddress','hostname','image_id','stack','name','state')]
            [set_key_b(d2,j2,m) for m in ('id','network','name','resources','parent','instanceId','privateAddress','hostname','image_id','stack','state','template')]
            """
            set_key(d2,j2,'name')
            set_key(d2,j2,'resources','network','name') --> "VLAN"
            set_key(d2,j2,'privateAddress')   --> "Machine"
            set_key(d2,j2,'hostnamename')
            set_key(d2,j2,'image_id ')
            set_key(d2,j2,'stack','name')
            set_key(d2,j2,'state') --> "Status"
            set_key(d2,j2,'template','name')  --> "Component"
            """
            #print(d2)
            l3 = []
            #d2.setdefault('disks',l3)
        

            if j2.has_key('resources'):
                for j in j2['resources']:
                    if j.has_key('resourceType'):
                        if j['resourceType'] == 'DISK_DRIVE':
                            l3.append((int(j['quantity'])))


            
            if len(l3) > 0:
                d2.setdefault('disks',','.join([str(i) for i in l3]))
            else:
                d2.setdefault('disks',None)
            
            pp(d2['disks'])
             
            disk_total = 0
            

            #pp(d2['disks_total'])
            if not d2['disks'] is None:
                for m in d2['disks'].split(','):
                    disk_total += int(m)
            d2.setdefault('disks_total',str(int(disk_total)))

            l.append(d2)
        return l

    else:
        print 'problem listing items for asset type %s'%asset_type
        print r.status_code, url
        return None

def create_xls_compute(l,site,cmp_cfg):
    wb2 = xlwt.Workbook(encoding='utf-8')

    ts = datestamp5()
    dp_root = os.path.join('../reports/compute/xls',ts)
    print 'creating report(s) in %s'%dp_root
    try:
        os.mkdir(dp_root)
    except OSError:
        pass

    fp = os.path.join(dp_root,'%s-compute-%s.xls'%(site,ts))


    sht = wb2.add_sheet('index')

    #hdr = [i for i in l[0].keys() if i in ('hostname','id','image_id','instanceId','name','parent','privateAddress','stack','state')]
    #orig: hdr = [['hostname',6000],['id',2500],['image_id',2500],['instanceId',10000],['name',15000],['parent',6000],['privateAddress',6000],['stack',10000],['state',6000]]
    hdr      = [['hostname',6000],['name',15000],['Component',10000],['VLAN',6000],['Machine',6000],['stack',10000],['Status',6000]]
    hdr_keys = [['hostname',6000],['name',15000],['template',10000],['resources',10000],['privateAddress',6000],['stack',10000],['state',6000]]
    
    #pp(hdr)
    #pp(len(hdr))
    #pp(range(0,len(hdr)))
    hdr_style = xls.get_style(wrap=1,color=5,isBold=True,isBorders=True)
    #orig:hdr_style = xls.get_style(wrap=1,color=169,isBold=True,isBorders=True)
    xls.write_hdr(sht,hdr,0,hdr_style)
    xls.adjust_colwidths(sht,hdr)

    #for i in range(0,len(hdr)):
        #sht.write(0,i,hdr[i])

    nrow = 1
    style_yellow = xls.get_style(wrap=1,color=133,isBold=False,isBorders=False)
    style_green = xls.get_style(wrap=1,color=3,isBold=False,isBorders=False)
    for d in l:
        #pp(d)
        #for m,n in enumerate(hdr):
        for m,n in enumerate(hdr_keys):
            #print(m,n)	
            if n[0] == 'resources':
                 print('VLAN:',	d[n[0]] )
            if n[0] == 'stack':
                sht.write(nrow,m,d[n[0]]['name'])
            else:
               if n[0] != 'template':
                     if n[0] == 'state':
                              status_str = translate_machine_state(d[n[0]])
                              if status_str == 'ONLINE':
                                  sht.write(nrow,m,status_str,style_green)
                              else:
                                  if status_str == 'OFFLINE':
                                           sht.write(nrow,m,status_str,style_yellow)
                                  else:
                                           sht.write(nrow,m,status_str)
                     else:
                              sht.write(nrow,m,d[n[0]])
               else:
                    cid = d[n[0]]
                    sht.write(nrow,m, get_component_id_compute(cid,cmp_cfg))
        nrow += 1

            
    print "Saving reports %s",(fp)
    wb2.save(fp)
    return fp

def write_row(sht,r,vals,style): #cid,COMPID,parent_name,bpname,wkldname,pkgname,pkg_version,scrpt_type,scrpt_name,scrpt_version):
    for i,v in enumerate(vals):
        sht.write(r,i,v,style); 

def create_xls(l_wkld_pkg_pvars,d_pkgs,d_scrpts,bps_cfg,site):
    """
    Create xls report, requires "l4", from call to bps.rpt3 function. 
    Refer to test runner xls_baseline_R3_test_runner(), in test_blueprints for more detailed usage.
    
    in: l_wkld_pkg_pvars (bp,bp_vars,wkld,wkld_pvars,l_pkg_items):
            [   [BP-1351, [], WKLD-1352, [],
                [(PKG-67, [
                     ('PKG-67', u'MWS2R2 - AD Rename and Add to Domain', u'COMPONENTID', '', u'string-any', True)],
                    [('SCRPT-228', u'MWS2R2 - AD Global Repository', u'ENVIRONMENT', '', u'string-any', True),...],
                    [('I', SCRPT-228), ('I', SCRPT-229), ('I', SCRPT-230)]),
                   (PKG-52,
                    ...
    in: d_pkgs: {'asset_details': (PKG-26,), 'parent_details': [(BP-65, WKLD-62), (BP-61, WKLD-62)]}
        d_scrpts: {'asset_details': (SCRPT-128, 'I'), 'parent_details': [(BP-390, WKLD-417, PKG-30), (BP-171, WKLD-267, PKG-30)]}

    in: bps_cfg - only used to match component ID with a blueprint (bps_cfg[str(bp.id)]['CID']); a hack, should be a member of asset class
                or directly come from asset rather than external list here.

    """

    """
    def dump_bpvar_details(sht,nrow,bp,bp_vars,wkld,wkld_pvars,l_pkg_items):
        def write_row(sht,r,type_column_offset,cid,compnt_id,parent_name,bpname,asset_name,id,k,v,is_default):
                sht.write(r,0,cid) sht.write(r,1,compnt_id) #bps_cfg[str(bp.id)]['CID']) sht.write(r,2,parent_name) sht.write(r,3,bpname) sht.write(r,type_column_offset+4,id) sht.write(r,8,asset_name) sht.write(r,9,k) sht.write(r,10,v) sht.write(r,11,is_default)

        for var_id,s_name,k,v,ptype,is_default in sorted(bp_vars,key=lambda x:x[1]):
            type_column_offset = 0
            write_row(sht,nrow,type_column_offset,bp.short_id,compnt_id,bp.parent_name,bp.name,bp.name,var_id,k,v,is_default)
            nrow += 1
        return nrow
    """

    def dump_wkld_pkg_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id):

        isUnitTested = False
        for p,g in groupby(sorted(l_pkg_items,key=lambda x:x[0].name)):
            pkgs = list(g)
            for rec in pkgs:
                p,l_pvars,s_svars,l_scripts = rec
                if p.name == 'MWS2R2 - Unit Testing': isUnitTested = True
                for scrpt_type,s in l_scripts:
                    p_ver, s_ver = map_ver(p.version), map_ver(s.version)
                    write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,p.name,p_ver,scrpt_type,s.name,s_ver),style)
                    nrow += 1
        if not isUnitTested:
            write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,'*** NO UNIT TEST PACKAGE ***','***','***','***','***'),style_yellow)
            nrow += 1
        return nrow

    def dump_wkld_resource_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id,style,style_yellow):

        #pp(wkld_pvars)
        """
[('WKLD-266', u'DC12001', u'INSTANCEID', u'001', u'string-any', False),
 ('WKLD-266', u'DC12001', u'COMPONENTID', u'DC12', u'string-any', False)]
        """


        #desc = '%s-%s-%s'%(comp_id,inst_id,<ANDREWS_NEWDESCVALUE>)
        

        #desc = '%s-%s-%s'%(comp_id,inst_id,<ANDREWS_NEWDESCVALUE>)
        
        print(wkld)
        if len(wkld.ro.resources.items()) == 0:
            write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,'','','',''),style_yellow)
            nrow += 1
        else:
            comp_id,inst_id = "",""

            for k,v in wkld.ro.resources.iteritems():
                print("k:",k)
                print("v:",v)
                resource_type,resource_name,resource_value = v

                #START OF A.AD BLOCK
                #pp(resource_name)
                
                desc = ""
                #comp_id,inst_id = "",""
                for wkld_id,wkld_name,var_name,var_value,var_type,b in wkld_pvars:
                    #print(resource_name, wkld_id, var_name, var_value)
                                    
                    if var_name == 'COMPONENTID': 
                        comp_id = var_value
#                        comp_id = resource_name
                        #print(var_name, comp_id)
                    if var_name == 'INSTANCEID': 
                        inst_id = var_value   
#                        inst_id = resource_name 
                        #print(var_name, comp_id)  
                    #print(comp_id, "-", inst_id, "-", wkld.description)

                    #desc = '%s-%s-%s'%(comp_id,inst_id,wkld.description)


                    """
                    if comp_id != "" and inst_id != "" and wkld.description != "":
                        desc = '%s-%s-%s'%(wkld.description) 
                        #print (desc)
                    else:
                        desc = wkld.description
                    """
                if comp_id != "" and inst_id != "": 
                    if wkld.description != "":
                        desc = '%s-%s-%s'%(comp_id,inst_id,wkld.description) 
                    else:
                        desc = '%s-%s-%s'%(comp_id,inst_id,"PLEASE COMPLETE DESCRIPTION") 
                        
                        #print (desc,wkld.description)
                #else:
                    #if wkld.description != "":
                        #print("not all set",comp_id,inst_id,wkld.description)  
                if desc != "":
                    write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,desc,resource_type,resource_name,resource_value),style)
                    nrow += 1
                    #print ("desc:",desc)
                else:
                    write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,wkld.description,resource_type,resource_name,resource_value),style)
                    nrow += 1
                    #print("workload:",wkld.description)
                    
                                         
                #if desc != "":
                 #   print (desc)
                #END OF A.AD BLOCK  
 
                #write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,desc,resource_type,resource_name,resource_value),style)
                #write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,wkld.name,wkld.description,resource_type,resource_name,resource_value),style)
                #nrow += 1


        return nrow

    def dump_wkld_var_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id):

        for wkld_id,s_name,k,v,ptype,is_default in sorted(wkld_pvars,key=lambda x:x[1]):
                BASELINES = get_baseline_versions(bp.id,bps_cfg)
                is_bp,is_wkld,is_pkg,is_scrtp = '','','',''
                write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,'',wkld_id,'','',wkld.name,k,v,is_default),style)
                nrow += 1

        for p,l_pvars,s_svars,l_scripts in sorted(l_pkg_items,key=lambda x:x[0].name):

            for pkg_id,s_name,k,v,ptype,is_default in sorted(l_pvars,key=lambda x:x[1]):
                    BASELINES = get_baseline_versions(bp.id,bps_cfg)
                    is_bp,is_wkld,is_pkg,is_scrtp = '','','',''
                    write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,'','',pkg_id,'',p.name,k,v,is_default),style)
                    nrow += 1

            for scrpt_id,s_name,k,v,ptype,is_default in sorted(list(s_svars),key=lambda x:x[1]):
                    BASELINES = get_baseline_versions(bp.id,bps_cfg)
                    is_bp,is_wkld,is_pkg,is_scrtp = '','','',''
                    write_row(sht,nrow,(bp.short_id,compnt_id,bp.parent_name,bp.name,'','','',scrpt_id,s_name,k,v,is_default),style)
                    nrow += 1

        return nrow

    def create_summ_sht(wb2,style):
        sht = wb2.add_sheet('Index')

    def create_index_sht(wb2,style,l_wkld_pkg_pvars):
        """
        Using l_wkld_pkg_pvars as the data source, however it is many workloads to one blueprint. This report is blueprint only; 
        so use groupby and take the first record and report on this one.
        """
        hdr = [('bp_id',4000),('component',4000),('bp_details',19000),('last_modified',4000),('container',8000),('creator',6000),('baselines',6000)]
        sht = wb2.add_sheet('BP_INDEX')
        nrow = 0
        sht.write(nrow,0,'=HYPERLINK("#%s!A1","%s")'%('BP_VARS_ALL','All Variables'))
        nrow = 1
        sht.write(nrow,0,'Blueprint Index',style)
        nrow = 2
        xls.write_hdr(sht,hdr,nrow,hdr_style)
        xls.adjust_colwidths(sht,hdr)
        # write index
        nrow, ncol = 3,0
        l2 = []
        for bp,g in groupby(sorted(l_wkld_pkg_pvars,key=lambda x:x[0]),key=lambda y:y[0]):
            l2.append(list(g)[0])

        for rec in sorted(l2,key=lambda x:x[-1]):
            bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = rec
            BASELINES = get_baseline_versions(bp.id,bps_cfg)
            write_row(sht,nrow,(bp.id,compnt_id, '=HYPERLINK("#%s!A1","%s")'%(bp.id,bp.name),timestamp_to_date(bp.mtime),bp.parent_name,bp.creator,BASELINES),style)
            nrow += 1

    def create_resource_sht(wb2,style,l_wkld_pkg_pvars,d_pkgs,d_scrpts):
        hdr = [('bp_id',2500),('component',3000),('bp_container',4000),('bp_name',10000),('wkld_name',6000),('wkld_desc',12000),('resource type',5500),('resource name',7000),('resource value',3000)]
        sht = wb2.add_sheet('BP_WKLD_RESOURCES')
        nrow = 1
        xls.write_hdr(sht,hdr,nrow,hdr_style)
        nrow, ncol = 2,0
        for bp,g in groupby(sorted(l_wkld_pkg_pvars,key=lambda x:x[-1]),key=lambda y:y[-1]):
            wklds = list(g)
            for rec in wklds: 
                bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = rec
                nrow = dump_wkld_resource_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id,style,style_yellow)
        xls.adjust_colwidths(sht,hdr)

    def create_pkg_sht(wb2,style,l_wkld_pkg_pvars,d_pkgs,d_scrpts):
        hdr = [('bp_id',2500),('component',3000),('bp_container',4000),('bp_name',12000),('wkld_name',6000),('pkg_name',11000),('pkg version',4000),('script type',3000),('script_name',11000),('script version',4000)]
        sht = wb2.add_sheet('BP_PKG_SCRPT_VER')
        nrow = 1
        xls.write_hdr(sht,hdr,nrow,hdr_style)
        nrow, ncol = 2,0
        for bp,g in groupby(sorted(l_wkld_pkg_pvars,key=lambda x:x[-1]),key=lambda y:y[-1]):
            wklds = list(g)
            for rec in wklds: 
                bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = rec
                nrow = dump_wkld_pkg_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id)
        xls.adjust_colwidths(sht,hdr)

    def create_var_sht(wb2,style,l_wkld_pkg_pvars):
        """
        l_wkld_pkg_pvars: ([[BP-309, [('BP-309', u'SCID', u'XDM'), ('BP-309', u'CID', u'EM')], WKLD-310, [(PKG-67, [('PKG-67', u'INSTANCEID', ''), ('PKG-67
        """
        hdr = [('bp_id',2500),('component',4000),('bp_container',6000),('bp_name',10000),('is_bp',2800),('is_wkld',2800),('is_pkg',2800),('is_scrpt',2800),('asset_name',11000),('var_name',6000),('var_value',16000),('isdefault',2500)]
        sht = wb2.add_sheet('BP_VARS_ALL')
        sht.write(0,0,'Blueprint Variables',style)
        nrow = 1
        xls.write_hdr(sht,hdr,nrow,hdr_style)
        nrow, ncol = 2,0
        for bp0,g0 in groupby(sorted(l_wkld_pkg_pvars,key=lambda m:m[-1]),key=lambda n:n[-1]):
            for bp,g in groupby(sorted(list(g0),key=lambda x:x[0]),key=lambda y:y[0]):
                wklds = list(g)

                for rec in wklds: 
                    bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = rec
                    nrow = dump_wkld_var_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id)

        xls.adjust_colwidths(sht,hdr)

    def create_detail_sht(wb2,style,wklds):
        hdr = [('bp_id',4500),('component',4000),('bp_container',6000),('bp_name',10000),('is_bp',2800),('is_wkld',2800),('is_pkg',2800),('is_scrpt',2800),('asset_name',11000),('var_name',6000),('var_value',16000),('isdefault',2500)]

        bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = wklds[0]
        sht = wb2.add_sheet(str(bp.id))
        sht.col(0).width = 26000
        sht.col(1).width = 26000

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        nrow,ncol = 0,0
        s = 'Blueprint configuration sheet for %s'%bp.name
        sht.write(nrow,ncol,s,style)

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        nrow += 2; ncol = 0

        def l1(k,v,nrow,ncol):
            sht.write(nrow,ncol,k,style)
            sht.write(nrow,ncol+1,v,style)

        nrow +=1; l1('Name',bp.name,nrow,0)
        nrow +=1; l1('Description',bp.description,nrow,0)
        nrow +=1; l1('Creator',bp.creator,nrow,0)
        nrow +=1; l1('Last Modified',timestamp2date(bp.mtime),nrow,0)
        nrow +=1; l1('Environment',bp.parent_name,nrow,0)

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        nrow += 2; ncol = 0
        sht.write(nrow,ncol,'Variable Details', style)
        nrow += 1

        xls.write_hdr(sht,hdr,nrow,hdr_style)
        nrow += 1
        for rec in wklds:
            bp,bp_vars,wkld,wkld_pvars,l_pkg_items,compnt_id = rec
            nrow = dump_wkld_var_details(sht,nrow,bp,wkld,wkld_pvars,l_pkg_items,compnt_id)

        xls.adjust_colwidths(sht,hdr)

    hdr_style = xls.get_style(wrap=1,color=169,isBold=True,isBorders=True)
    style_yellow = xls.get_style(wrap=1,color=133,isBold=False,isBorders=False)
 
    style = xlwt.XFStyle()
    style.alignment.wrap = 1

    wb2 = xlwt.Workbook(encoding='utf-8')

    ts = datestamp5()
    dp_root = os.path.join('../reports/baseline/xls',ts)
    print 'creating report(s) in %s'%dp_root
    try:
        os.mkdir(dp_root)
    except OSError:
        pass

    #fp = os.path.join(dp_root,'blueprints-%s.xls'%ts)
    fp = os.path.join(dp_root,'%s-blueprints-%s.xls'%(site,ts))

    def update_l_wkld_pkg_pvars(l_wkld_pkg_pvars):
        """
            appends component ID to each record and then sorts l_wkld_pkg_pvars on it. rec[0] is first element of l_wkld_pkg_pvars
            list, eg BP-878:
            
            l_wkld_pkg_pvars:
                [BP-878,
                 [],
                 WKLD-883,
                 [('WKLD-883', u'APP003', u'COMPONENTID', u'APP', u'string-any', False),
                  ('WKLD-883',
                <NEW COMPONENT ID>
                ]

        """
        [rec.append(get_component_id(rec[0].id,bps_cfg)) for rec in l_wkld_pkg_pvars]
#        return sorted(l_wkld_pkg_pvars,key=lambda x:x[-1])

    update_l_wkld_pkg_pvars(l_wkld_pkg_pvars)

    create_summ_sht(wb2,style)#:sorted(l5))
    create_index_sht(wb2,style,l_wkld_pkg_pvars)#:sorted(l5))
    create_resource_sht(wb2,style,l_wkld_pkg_pvars,d_pkgs,d_scrpts)
    create_pkg_sht(wb2,style,l_wkld_pkg_pvars,d_pkgs,d_scrpts)
    create_var_sht(wb2,style,l_wkld_pkg_pvars)

    for bp,g in groupby(sorted(l_wkld_pkg_pvars,key=lambda x:x[0].id),key=lambda y:y[0].id):
        wklds = list(g)
        create_detail_sht(wb2,style,wklds)

    print "Saving reports %s",(fp)
    wb2.save(fp)
    return fp

def create_docs(l4):
    """
    in:
 (BP-CID-SCID-983-MWS2R2___OWA_2013,
  [(u'CERTIFICATE_NAME', ''),
   (u'PIDKEY', ''),
   (u'INSTALL_MEDIA', u'InstallMedia_OWA'),
   (u'INTERNAL_URL', ''),
   (u'INSTALLLOCATION', ''),
   (u'EXTERNAL_URL', '')],
  [(WKLD-CID-SCID-989-APP005,
    {u'COMPONENTID': u'APP',
     u'INSTANCEID': u'005',
     u'OWAHOST_COMPONENTID': u'APP',
     u'OWAHOST_INSTANCEID': u'005'},
    [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
     PKG-CID-SCID-238-MWS2R2___OWA_2013_Deploy_Install_Files,
     PKG-CID-SCID-240-MWS2R2___OWA_2013_Install,
     PKG-CID-SCID-241-MWS2R2___OWA_2013_Configure]),
   (WKLD-CID-SCID-991-APP006,
    {u'COMPONENTID': u'APP',
     u'INSTANCEID': u'006',
     u'OWAHOST_COMPONENTID': u'APP',
     u'OWAHOST_INSTANCEID': u'005'},
    [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
     PKG-CID-SCID-238-MWS2R2___OWA_2013_Deploy_Install_Files,
     PKG-CID-SCID-240-MWS2R2___OWA_2013_Install,
     PKG-CID-SCID-241-MWS2R2___OWA_2013_Configure])]),

    """

    def create_doc(fp,bp,bp_diff_vars,wklds):

        document = Document()

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_heading(bp.short_id, 0)
        p = document.add_paragraph('Blueprint build document for ')
        p.add_run(' %s '%bp.name).bold = True



        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        # For ToC

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        document.add_heading('Blueprint Information', level=1)
        document.add_paragraph()

        def write_table1(recordset,ncols):
            table = document.add_table(rows=0, cols=ncols,style="TableGrid")
            for k in recordset:
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
            document.add_paragraph()
   
        def write_table2(recordset,ncols):
            table = document.add_table(rows=0, cols=ncols,style="TableGrid")
            for k,v in recordset:
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                row_cells[1].text = str(v)
            document.add_paragraph()

        def l1(k,v):
            p = document.add_paragraph()
            p.add_run('%s: '%k).bold = True
            p.add_run(v)

        l1('Name',bp.name)
        l1('Description',bp.description)
        l1('Creator',bp.creator)
        l1('Last Modified',timestamp2date(bp.mtime))
        l1('Environment',bp.parent_name)


        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        document.add_heading('Workload Details', level=1)

        document.add_paragraph()
    
        """
   PKG-CID-SCID-265-MWS2R2___MCMSZC_Install_msi])]
(WKLD-CID-SCID-1025-MCM001_Windows_2012_R2_S___,
 {u'COMPONENTID': u'MCM', u'INSTANCEID': u'001', u'SOURCE': u'MCMSZC'},
 [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
  PKG-CID-SCID-262-MWS2R2___Get_App_Files,
  PKG-CID-SCID-265-MWS2R2___MCMSZC_Install_msi])
        """
        def write_wkld(wk,wk_vars,wk_pkgs):
            document.add_heading(wk.name, level=2)
            l1('Packages','\n'+'\n'.join([p.short_id for p in wk_pkgs]))

        for wk,wk_vars,wk_pkgs in wklds:
            write_wkld(wk,wk_vars,wk_pkgs)


        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        document.add_heading('Variable Details', level=1)
        document.add_paragraph()

        document.add_heading('Blueprint-wide (Script Global) Variables', level=2)
        document.add_paragraph('These values are not set in any workloads, any default values come from what is set globally within a script. Refer to section Script Default Variables and Values.')

        write_table2(sorted(bp_diff_vars,key=lambda x:x[0]),ncols=2)

        document.add_heading('Workload Variables', level=2)
        for wk,wk_vars,wk_pkgs in wklds:
            l1('Workload',wk.name)
            write_table2(sorted([i for i in wk_vars.iteritems()],key=lambda x:x[0]),ncols=2)

        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        document.add_heading('Package-Script Cross Reference', level=1)

        document.add_paragraph('Note: I = Installer script; O = Operational script. If script is both installer and operational within same package, the operational script is not listed.')

        d2 = {}
        for wk,wk_vars,wk_pkgs in wklds:
            for p in wk_pkgs:
                    d2.setdefault(p.short_id,{})
                    d2[p.short_id]['package'] = p
                    d2[p.short_id].setdefault('scripts',set())
                    for j in p.scripts:
                        d2[p.short_id]['scripts'].add(j)

        """
        d2:

 u'PKG-CID-SCID-262-MWS2R2___Get_App_Files': {'package': PKG-CID-SCID-262-MWS2R2___Get_App_Files,
                                              'scripts': set([('I',
                                                               SCRIPT-CID-SCID-681-MWS2R2___Get_App_Files)])},
 u'PKG-CID-SCID-265-MWS2R2___MCMSZC_Install_msi': {'package': PKG-CID-SCID-265-MWS2R2___MCMSZC_Install_msi,
                                                   'scripts': set([('I',
                                                                    SCRIPT-CID-SCID-683-MWS2R2___SZC_Install_msi)])}}

        """

        
        for k,v in sorted(d2.iteritems()):
            document.add_heading(k, level=2)
            document.add_heading('Scripts', level=3)
            write_table2(sorted(list(v['scripts'])),ncols=2)


        #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
        document.add_page_break()
        document.add_heading('Script Details', level=1)

        document.add_paragraph('Note that script varibles are global across solution unless specifically set against a blueprint, workload or package. Refer to section Blueprint-Wide variables.')

        document.add_paragraph('Refer to Blueprint build document MWS2R2 Active Directory for scripts: \nMWS2R2 - AD Rename Computer\nMWS2R2 - AD Global Repository\nMWS2R2 - AD Turn Off Windows Firewall\nMWS2R2 - AD Add To Domain')

        def f4(p):
            if bp.name == 'MWS2R2 Active Directory':
                return True  
            elif p.name in ['MWS2R2 - AD Rename and Add to Domain','MWS2R2 - AD Rename Computer','MWS2R2 - AD Global Repository','MWS2R2 - AD Turn Off Windows Firewall','MWS2R2 - AD Add To Domain']:
                return False
            else:
                return True

        s2 = set()
        for wk,wk_vars,wk_pkgs in wklds:
            for p in filter(f4,wk_pkgs):
                    for s in p.scripts:
                        s2.add(s[1])


        for s in s2:
            #~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~
            document.add_page_break()
            document.add_heading(s.name, level=2)

            document.add_heading('Default Variables and Values', level=3)
            write_table1(sorted(s.vars,key=lambda x:x[0]),ncols=2)

            document.add_heading('Body', level=3)
            l1('Body',s.body) 
            
        document.save(fp)
    
    dp_root = os.path.join('../reports/baseline/docx',datestamp5())
    print 'creating report(s) in %s'%dp_root
    os.mkdir(dp_root)

    for bp,bp_diff_vars,wklds in l4:
        tdp = os.path.join(dp_root,bp.parent_name)
        try:
            os.mkdir(tdp)
        except:
            pass
        fp = os.path.join(tdp,'%s.docx'%bp)
        create_doc(fp,bp,bp_diff_vars,wklds)

def bps_rpt(bp_list):
    """
        in (l4):
         (BP-CID-SCID-983-MWS2R2___OWA_2013,
          [(u'CERTIFICATE_NAME', ''),
           (u'PIDKEY', ''),
           (u'INSTALL_MEDIA', u'InstallMedia_OWA'),
           (u'INTERNAL_URL', ''),
           (u'INSTALLLOCATION', ''),
           (u'EXTERNAL_URL', '')],
          [(WKLD-CID-SCID-989-APP005,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'005',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'}),
           (WKLD-CID-SCID-991-APP006,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'006',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'})]),

        out:
            (u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             '20150419',
             u'MWS2_CDCR2',
             u'mkrynina',
             [u'WKLD-CID-SCID-845-APP005', u'WKLD-CID-SCID-1374-APP006'],
             BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm')
    """
    l = []
    for bp in bp_list:
        l.append((bp.short_id,timestamp2date(bp.mtime),bp.parent_name,bp.creator,[w.short_id for w in bp.workloads],bp))
    return l


def wklds_rpt(l4):
    """
        out:
            (u'WKLD-CID-SCID-845-APP005',
             '20150419',
             u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             [(u'PKG-CID-SCID-67-MWS2R2___AD_Rename_and_Add_to_Domain',
               {u'COMPONENTID': u'LNC', u'INSTANCEID': u'001'}),
              (u'PKG-CID-SCID-52-MWS2R2___Set_PageFile', {}),
              (u'PKG-CID-SCID-53-MWS2R2___Set_Volume_Label', {}),
              (u'PKG-CID-SCID-37-MWS2R2___OWA_2013_Install',
               {u'INSTALL_MEDIA': u'InstallMedia', u'SOURCE': u'OfficeWebApps2013'}),
              (u'PKG-CID-SCID-38-MWS2R2___OWA_2013_Configure', {})])
    """
    l = []
    for bp,bp_diff_vars,wklds in l4:
        for w in wklds:
            cid,bp_cid,mtime,wkld = w[0].short_id,w[0].bp_cid,w[0].mtime,w[0]
            l.append((cid,timestamp2date(mtime),bp_cid,[(j.short_id,j.vars) for j in wkld.packages],w[0]))
    return l

def scrpts_rpt(l4):
    """
        in (l4):
         (BP-CID-SCID-983-MWS2R2___OWA_2013,
          [(u'CERTIFICATE_NAME', ''),
           (u'PIDKEY', ''),
           (u'INSTALL_MEDIA', u'InstallMedia_OWA'),
           (u'INTERNAL_URL', ''),
           (u'INSTALLLOCATION', ''),
           (u'EXTERNAL_URL', '')],
          [(WKLD-CID-SCID-989-APP005,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'005',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'}),
           (WKLD-CID-SCID-991-APP006,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'006',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'})]),

        out:
            (u'SCRIPT-CID-SCID-228-MWS2R2___AD_Global_Repository',
             '20150416',
             u'PKG-CID-SCID-67-MWS2R2___AD_Rename_and_Add_to_Domain',
             u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             'I',
             {u'ENVIRONMENT': ''},
             SCRIPT-CID-SCID-228-MWS2R2___AD_Global_Repository)
    """
    l = []
    d_scrpt_bps = dict()
    for bp,bp_diff_vars,wklds in l4:
        for w in wklds:
            cid,bp_cid,mtime,wkld = w[0].short_id,w[0].bp_cid,w[0].mtime,w[0]
            for p in wkld.packages:
                for script_type,s in p.scripts:
                    d_scrpt_bps.setdefault(s.short_id,[])
                    d_scrpt_bps[s.short_id].append(bp_cid)
                    if s.short_id not in [i[0] for i in l]:
                        l.append([s.short_id,timestamp2date(s.mtime),p.short_id,None,script_type,s.vars,s.name,s])
    for s in l:
        s[3] = d_scrpt_bps[s[0]]
    return sorted(l,key=lambda x:x[-2])

         

import xlwt
def baseline_rpt():
    """
        in:


    """

#    asset_names = ['BP-CID-SCID-1161-MWS2R2_SQLServer_Cluster_SharePoint','BP-CID-SCID-1192-MWS2R2_SQLServer_Cluster_Lync']
    asset_names = None
    #asset_names = ['BP-CID-SCID-1161-MWS2R2_SQLServer_Cluster_SharePoint']
    #asset_names = ['BP-CID-SCID-983-MWS2R2___OWA_2013']
    asset_names = ['BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm']

    bps.load_assets(asset_types=['blueprint'],asset_names=asset_names,maxn=999)
          
    wb2 = xlwt.Workbook(encoding='utf-8')

    style = xlwt.XFStyle()
    style.alignment.wrap = 1

    summ_sht = wb2.add_sheet('Summary')

    ######## bp_rpt_sht ######
    bp_rpt_sht = wb2.add_sheet('Blueprints')
    """
    bp_rpt = bps_rpt(bp_list)
        out:
            (u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             '20150419',
             u'MWS2_CDCR2',
             u'mkrynina',
             [u'WKLD-CID-SCID-845-APP005', u'WKLD-CID-SCID-1374-APP006'],
             BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm')

    hdr = ('bp_cid','bp_mtime','bp_container','bp_creator','bp_workloads')
    nrow,ncol = 0,0
    for j in hdr:
        bp_rpt_sht.write(nrow,ncol,j)
        ncol += 1

    for i in bp_rpt:
        ncol = 0
        nrow += 1
        for j in range(0,len(hdr)-1):
            bp_rpt_sht.write(nrow,ncol,i[j],style)
            ncol += 1
        bp = i[-1]
        bp_rpt_sht.write(nrow,ncol,'\n'.join([j.short_id for j in bp.workloads]),style)
    """

    bp_list = get_assets('blueprint')
    """
        in (l4):
         (BP-CID-SCID-983-MWS2R2___OWA_2013,
          [(u'CERTIFICATE_NAME', ''),
           (u'PIDKEY', ''),
           (u'INSTALL_MEDIA', u'InstallMedia_OWA'),
           (u'INTERNAL_URL', ''),
           (u'INSTALLLOCATION', ''),
           (u'EXTERNAL_URL', '')],
          [(WKLD-CID-SCID-989-APP005,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'005',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'}),
           (WKLD-CID-SCID-991-APP006,
            {u'COMPONENTID': u'APP',
             u'INSTANCEID': u'006',
             u'OWAHOST_COMPONENTID': u'APP',
             u'OWAHOST_INSTANCEID': u'005'})]),
    """

    hdr = ('bp_cid','bp_mtime','bp_container','bp_creator','bp_workloads','bp_vars','wkld_vars')
    nrow,ncol = 0,0
    for j in hdr:
        bp_rpt_sht.write(nrow,ncol,j)
        ncol += 1

    for i in bp_list:
        nrow += 1
        ncol = 0
        pp(i)
        bp_rpt_sht.write(nrow,ncol,i[0],style)

    bp_rpt_sht.col(0).width = 11500
    bp_rpt_sht.col(1).width = 3000
    bp_rpt_sht.col(2).width = 8500
    bp_rpt_sht.col(3).width = 6000
    bp_rpt_sht.col(4).width = 13000
    bp_rpt_sht.col(5).width = 6000

    ######## wkld_rpt_sht ######
    wkld_rpt_sht = wb2.add_sheet('Workloads')
    wkld_rpt = wklds_rpt(rpt3(bp_list))
    """
        out:
            (u'WKLD-CID-SCID-845-APP005',
             '20150419',
             u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             [(u'PKG-CID-SCID-67-MWS2R2___AD_Rename_and_Add_to_Domain',
               {u'COMPONENTID': u'LNC', u'INSTANCEID': u'001'}),
              (u'PKG-CID-SCID-52-MWS2R2___Set_PageFile', {}),
              (u'PKG-CID-SCID-53-MWS2R2___Set_Volume_Label', {}),
              (u'PKG-CID-SCID-37-MWS2R2___OWA_2013_Install',
               {u'INSTALL_MEDIA': u'InstallMedia', u'SOURCE': u'OfficeWebApps2013'}),
              (u'PKG-CID-SCID-38-MWS2R2___OWA_2013_Configure', {})])
    """
    hdr = ('wk_cid','wk_mtime','blueprint','blueprint_container','wk_vars','wk_pkgs','wk_pkg_vars')
    nrow,ncol = 0,0
    for j in hdr:
        wkld_rpt_sht.write(nrow,ncol,j)
        ncol += 1

    def f2(p,p_vars):
        return p+'\n'+'\n\t'.join(['%s:%s'%(j,k) for j,k in p_vars.iteritems()])

    def f3(wk_vars):
        return '\n\t'.join(['%s:%s'%(j,k) for j,k in wk_vars.iteritems()])


    for i in wkld_rpt:
        ncol = 0
        nrow += 1
        for j in range(0,len(hdr)-4):
            wkld_rpt_sht.write(nrow,ncol,str(i[j]),style)
            ncol += 1
        w = i[-1]
        parent_name,wk_vars = w.parent_name,w.vars
        pp(parent_name)
        pp(wk_vars)
        wkld_rpt_sht.write(nrow,ncol,parent_name,style)
        ncol += 1
        wkld_rpt_sht.write(nrow,ncol,f3(wk_vars),style)
        ncol += 1
        pkgs = i[-2]
        wkld_rpt_sht.write(nrow,ncol,'\n'.join([p for p,p_vars in pkgs]),style)
        ncol += 1
        wkld_rpt_sht.write(nrow,ncol,'\n'.join([f2(p,p_vars) for p,p_vars in pkgs]),style)

    wkld_rpt_sht.col(0).width = 9500
    wkld_rpt_sht.col(1).width = 3000
    wkld_rpt_sht.col(2).width = 10500
    wkld_rpt_sht.col(3).width = 4000
    wkld_rpt_sht.col(4).width = 16000
    wkld_rpt_sht.col(5).width = 15000
    wkld_rpt_sht.col(6).width = 15000


    ######## scrpt_rpt_sht ######
    scrpt_rpt_sht = wb2.add_sheet('Scripts')
    scrpt_rpt = scrpts_rpt(rpt3(bp_list))
    """
        out:
            (u'SCRIPT-CID-SCID-228-MWS2R2___AD_Global_Repository',
             '20150416',
             u'PKG-CID-SCID-67-MWS2R2___AD_Rename_and_Add_to_Domain',
             u'BP-CID-SCID-839-MWS2R2___Office_Web_Apps_2013_Farm',
             'I',
             {u'ENVIRONMENT': ''},
             'AD_Global_Repository',
             SCRIPT-CID-SCID-228-MWS2R2___AD_Global_Repository)
    """
    hdr = ('scrpt_cid','scrpt_mtime','scrpt_package','scrpt_blueprint','script_type','script_vars','script')
    nrow,ncol = 0,0
    for j in hdr:
        scrpt_rpt_sht.write(nrow,ncol,j)
        ncol += 1

    def f3(s_vars):
        return '\n'.join(['%s:%s'%(j,k) for j,k in s_vars.iteritems()])

    for i in scrpt_rpt:
        ncol = 0
        nrow += 1
        for j in range(0,len(hdr)-2):
            scrpt_rpt_sht.write(nrow,ncol,i[j],style)
            ncol += 1
        scrpt_vars = i[-3]
#        scrpt_rpt_sht.write(nrow,ncol,'\n'.join(scrpt_vars.keys()),style)
        scrpt_rpt_sht.write(nrow,ncol,f3(scrpt_vars),style)
        ncol += 1
        scrpt_rpt_sht.write(nrow,ncol,i[-1].body,style)

    scrpt_rpt_sht.col(0).width = 10500
    scrpt_rpt_sht.col(1).width = 3000
    scrpt_rpt_sht.col(2).width = 10500
    scrpt_rpt_sht.col(3).width = 10500
    scrpt_rpt_sht.col(4).width = 3000
    scrpt_rpt_sht.col(5).width = 18000

    dp_root = os.path.join('../reports/baseline',datestamp5())
    os.makedirs(dp_root)

    dp_rpt = os.path.join(dp_root,'baseline.xls')
    print "Saving reports %s",(dp_rpt)
    wb2.save(dp_rpt)

def rpt_wkld_resources(b):
    return [(b[m],item.resources,) for m,sublist in [(k,i.workloads) for k,i in enumerate(b)] for item in sublist]

def rpt_wkld_vars(b):
    return [(b[m],item.vars,) for m,sublist in [(k,i.workloads) for k,i in enumerate(b)] for item in sublist]

def rpt_wkld_pkgs(b):
    """
    Generate packages report, top-down by blueprint, indexed by blueprint

in (b):
    b = bps.get_assets('BP')

    called: l_wkld_pkgs = bps.rpt_wkld_pkgs(b)   

out (l_wkld_pkgs):

(BP-CID-SCID-1160-MWS2R2___MTC,
 WKLD-CID-SCID-1161-MTC002_Windows_2012_R2_S___,
 [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
  PKG-CID-SCID-262-MWS2R2___Get_App_Files]),
(BP-CID-SCID-1160-MWS2R2___MTC,
 WKLD-CID-SCID-1162-MTC001_Windows_2012_R2_S___,
 [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
  PKG-CID-SCID-262-MWS2R2___Get_App_Files])]
    """

    return [(b[m],item,item.packages,) for m,sublist in [(k,i.workloads) for k,i in enumerate(b)] for item in sublist]

def rpt_wkld_pkg_pvars(l_wkld_pkgs):
    """
    Generate variables report, top-down by blueprint, but indexed by workload
    

    called: l_wkld_pkg_pvars = rpt_wkld_pkg_pvars(l_wkld_pkgs)

    returns l_wkld_pkg_pvars [[BP,BP_VARS,WKLD,[PKG,PVARS,SVARS]]]

out:
[   [BP-1351,
     [],
     WKLD-1352,
     [],
     [(PKG-32,
       [('PKG-32', u'MWS2R2 - AD Rename and Add to Domain', u'INSTANCEID', u'001', u'string-any', True),
        ('PKG-32', u'MWS2R2 - AD Rename and Add to Domain', u'COMPONENTID', u'LNC', u'string-any', True)],
       [('SCRPT-132', u'MWS2R2 - AD Global Repository', u'ENVIRONMENT', '', u'string-any', True),
        ('SCRPT-134', u'MWS2R2 - AD Rename Computer', u'COMPONENTID', '', u'string-any', True),
        ('SCRPT-134', u'MWS2R2 - AD Rename Computer', u'INSTANCEID', '', u'string-any', True)])]],
    [BP-1352,
     [],
     WKLD-1352,
     [],
     [(PKG-32,...
]

    """
    l2 = []
    for bp,wkld,pkgs in l_wkld_pkgs:
        l3 = []
        for p in pkgs:
            l3.append((p,p.vars,p.s_vars,p.scripts))
        l4 = [bp,bp.vars,wkld,wkld.vars]
        l4.append(l3)
        l2.append(l4)
    return l2

def rpt_bp_scrpts(l_wkld_pkgs):
    """
    return dictionary of scripts, top down from blueprint

in (l_wkld_pkgs):

(BP-CID-SCID-1160-MWS2R2___MTC,
 WKLD-CID-SCID-1161-MTC002_Windows_2012_R2_S___,
 [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
  PKG-CID-SCID-262-MWS2R2___Get_App_Files]),
(BP-CID-SCID-1160-MWS2R2___MTC,
 WKLD-CID-SCID-1162-MTC001_Windows_2012_R2_S___,
 [PKG-CID-SCID-235-MWS2R2___AD_Rename_and_Add_to_Domain,
  PKG-CID-SCID-262-MWS2R2___Get_App_Files])]

>>> d_scrpts = rpt.rpt_bp_scrpts(l_wkld_pkgs)
>>> len(d_scrpts.keys())
67
>>> d_scrpts.keys()[0]
128
>>> pp(d_scrpts[128])
{'asset_details': (SCRPT-128, 'I'),
 'parent_details': [(BP-390, WKLD-417, PKG-30),
                    (BP-1153, WKLD-1154, PKG-30),
                    (BP-1153, WKLD-1227, PKG-30),
                    (BP-1153, WKLD-1228, PKG-30),
                    (BP-171, WKLD-267, PKG-30)]}


    """
    d = {}
    for bp,wkld,pkgs in l_wkld_pkgs:
        for p in pkgs:
            for s_type,s in p.scripts:
                d.setdefault(s.id,{'asset_details':(s,s_type),'parent_details':[]})
                d[s.id]['parent_details'].append((bp,wkld,p))
    return d

def rpt_bp_pkgs(l_wkld_pkgs):
    """
    return dictionary of pkgs, top down from blueprint

>>> d_pkgs = rpt.rpt_bp_pkgs(l_wkld_pkgs)
>>> d_pkgs.keys()[0]
26
>>> pp(d_pkgs[26])
{'asset_details': (PKG-26,),
 'parent_details': [(BP-65, WKLD-62), (BP-61, WKLD-62)]}

    """
    d = {}
    for bp,wkld,pkgs in l_wkld_pkgs:
        for p in pkgs:
            d.setdefault(p.id,{'asset_details':(p,),'parent_details':[]})
            d[p.id]['parent_details'].append((bp,wkld))
    return d

def rpt_bp_bps(l_wkld_pkgs):
    """
    return dictionary of blueprints, top down from blueprint


>>> d_bps = rpt.rpt_bp_bps(l_wkld_pkgs)
>>> d_bps.keys()[0]
1153
>>> pp(d_bps[1153])
{'asset_details': (BP-1153,), 'parent_details': [[], [], []]}

    """
    d = {}
    for bp,wkld,pkgs in l_wkld_pkgs:
        d.setdefault(bp.id,{'asset_details':(bp,),'parent_details':[]})
        d[bp.id]['parent_details'].append(list())
    return d

